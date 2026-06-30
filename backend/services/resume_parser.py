#from user it will process check and then pass the file to backend pipeline


import io
import magic
from typing import Tuple, Optional, Tuple

import pdfplumber #parse pdf
from docx import Document #parse doc
import PyPDF2 #to fallback 

from backend.utils.file_utils import(
    FileParsingError, 
    TextExtractionError, 
    FileUploadError, 
    log_error, 
    log_warning, 
    log_info, 
    with_fallback
)

#size of files
from backend.core.config import (
    MAX_FILE_SIZE_BYTES,
    MAX_FILE_SIZE_MB, 
    SUPPORTED_MIME_TYPES
)

#Error will come two custom error 
class FileParsingError(Exception): #cant extract the text
    pass

class FileValidationError(Exception): #the file should not bt in system
    pass

def validate_file(file_data:bytes, filename:str)->Tuple[bool, str, Optional[str]]:
    file_size_bytes = len(file_data)

    #is size greater then this
    if file_size_bytes > MAX_FILE_SIZE_BYTES:
        size_mb = file_size_bytes / (1024 * 1024)
        return False, (
            f'File size ({size_mb:.2f} MB) exceeds the maximum of {MAX_FILE_SIZE_MB} MB. '
            'Please upload a smaller file or compress your resume.'
        ), None
    
    #if size 0 of file
    if file_size_bytes==0:
        return False, 'uploade file is empty...please check the file you have uploaded and try again'
    

    # Check internally the file type for proper validation 
    try:
        mime_type=magic.from_buffer(file_data, mime=True)
    except Exception as e:
        return False, f"error deteminin the file type : {e}", None
    

    #check in config that file type
    if mime_type not in SUPPORTED_MIME_TYPES:
        supported=', '.join(SUPPORTED_MIME_TYPES.keys()).upper()
        return False, (
            f'Unsupported file type: {mime_type}. '
            f'Please upload one of: {supported}.'
        ), None
    
    

    return True, '', SUPPORTED_MIME_TYPES[mime_type]







#in resume two layer front VISIBLE LAYER and a SECOND layer ANNOTATION LAYER with hyperlinks on second page
#FOR SECOND LAYER TO EXTRACT HYPERLINKS:
#TO validate the real links provided by the user like git 

def _extract_pdf_hyperlinks(file_data: bytes) -> str:
    urls = []
    try:
        reader = PyPDF2.PdfReader(io.BytesIO(file_data))
        for page in reader.pages:
            if '/Annots' not in page:
                continue
            for annot_ref in page['/Annots']:
                try:
                    annot = annot_ref.get_object()
                    if annot.get('/Subtype') != '/Link':
                        continue
                    action = annot.get('/A', {})
                    uri = action.get('/URI', '')
                    if uri and isinstance(uri, (str, bytes)):
                        # PyPDF2 may return bytes for URI values
                        if isinstance(uri, bytes):
                            uri = uri.decode('utf-8', errors='ignore')
                        uri = uri.strip()
                        if uri.startswith('http'):
                            urls.append(uri)
                except Exception:
                    pass
    except Exception:
        pass
    return '\n'.join(urls)


def _extract_pdf_with_pdfplumber(file_data: bytes) -> str:
    text = ''
    with pdfplumber.open(io.BytesIO(file_data)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + '\n'


#if empty
    if not text.strip():
        raise TextExtractionError(
            'pdfplumber extracted no text',
            user_message='No text could be extracted from the PDF.'
        )
    
    hyperlinks = _extract_pdf_hyperlinks(file_data)
    if hyperlinks:
        text = text.strip() + '\n' + hyperlinks

    return text.strip()


def _extract_pdf_with_pypdf2(file_data: bytes) -> str:
    text = ''
    pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_data))
    for page in pdf_reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + '\n'

#if empty
    if not text.strip():
        raise TextExtractionError(
            'PyPDF2 extracted no text',
            user_message='No text could be extracted from the PDF.'
        )

    hyperlinks = _extract_pdf_hyperlinks(file_data)
    if hyperlinks:
        text = text.strip() + '\n' + hyperlinks

    return text.strip()


def extract_text_from_pdf(file_data: bytes) -> str:
    try: 
        result, used_fallback=with_fallback(
        _extract_pdf_with_pdfplumber, 
        _extract_pdf_with_pypdf2, #if pdfpl not work then fallback to this
        file_data, 
        log_fallback=True
    )
    
        if used_fallback:
            log_info('PDF EXTRACTION succeded using the PyPDF2 fallback', context='resume_parser')
        return result
        
    except Exception as e:
        log_error(e, context='extract_text_from_pdf')
        raise FileParsingError(
            'Failed to extract text from PDF using both pdfplumber and PyPDF2. '
            'The PDF may be corrupted, password-protected, or contain only scanned images. '
            'Please ensure it contains selectable text.'
        ) from e
    


#For doc extraction
def extract_text_from_docx(file_data: bytes) -> str:
    try:
        doc = Document(io.BytesIO(file_data))
        text_parts = []


        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

#to extract from tables in doc is not done it might skip this part and ats score goes down
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text) #use append and strip because earlier work was in o(n) now will in big O(N)->FASTER

        text = '\n'.join(text_parts)

        if not text.strip():
            raise FileParsingError(
                'No text could be extracted from the document. '
                'The document may be empty or corrupted.'
            )
        
        try:
            for rel in doc.part.rels.values():
                if 'hyperlink' in rel.reltype.lower():
                    url = rel._target
                    if isinstance(url, str) and url.startswith('http'):
                        text += '\n' + url
        except Exception:
            pass

        log_info(f'Extracted {len(text)} chars from DOCX', context='resume_parser')
        return text.strip()

    except FileParsingError:
        raise   # Re-raise unchanged — don't wrap in another FileParsingError

    except Exception as e:
        log_error(e, context='extract_text_from_docx')
        raise FileParsingError(
            'Failed to extract text from DOCX. '
            'The document may be corrupted or in an unsupported format. '
            'Please try re-saving or converting to PDF.'
        ) from e

def extract_text_from_doc(file_data: bytes) -> str:
    raise FileParsingError(
        'Legacy .doc format is not supported. '
        'Please convert your document to .docx or .pdf and try again. '
        'You can convert using Microsoft Word, Google Docs, or online tools.'
    )


#ORCASTRATOR OF ALL FUNCTIONS :
def extract_text(file_data:bytes, file_type:str)->str:
    if file_type=='pdf':
        return extract_text_from_pdf(file_data)
    elif file_type=='docx':
        return extract_text_from_docx(file_data)
    elif file_type=='doc':
        return extract_text_from_doc(file_data)
    else:
        raise FileValidationError(
            f'invalid file type: {file_type}. supported types are: pdf, docx and doc'


        )
    




    #MASTER FUNCTION
def parse_resume_file(file_data: bytes, filename:str)->Tuple[str, dict]:
    log_info(f'parsing file :{filename}', context='parse_Resume_file')

    #phase01:validate file
    try:
        is_valid, error_msg, file_type=validate_file(file_data, filename)
        if not is_valid:
            log_warning(f'validation failed for file {filename}', context='parse_resume_file')
            raise FileValidationError(error_msg)
    
    except FileValidationError as e:
        raise 

    except Exception as e:
        log_error(e, context='parse_resume_file_validation')
        raise FileValidationError(
            'Could not validate the uploaded file. Please ensure it is a valid PDF or DOCX.'
        ) from e
    
    #phase02: extraction of file

    try:
        text = extract_text(file_data, file_type)
        log_info(f'Extracted {len(text)} chars from {filename}', context='parse_resume_file')

    except FileParsingError:
        raise   # Re-raise unchanged

    except Exception as e:
        log_error(e, context='parse_resume_file_extraction')
        raise FileParsingError(
            'An unexpected error occurred while processing the file. '
            'Please try again or contact support if the problem persists.'
        ) from e

    metadata = {
        'filename':        filename,
        'file_type':       file_type,
        'file_size_bytes': len(file_data),
        'text_length':     len(text),
        'success':         True,
    }
    return text, metadata